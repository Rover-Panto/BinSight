from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs"
WORK_DIR = ROOT / "work" / "binsight_assets"
LOGO_PNG = WORK_DIR / "monash_logo_white.png"
PDF_PATH = OUT_DIR / "BinSight_Final_Proposal.pdf"
DOCX_PATH = OUT_DIR / "BinSight_Final_Proposal.docx"

OUT_DIR.mkdir(exist_ok=True)

PAGE_W, PAGE_H = A4

NAVY = colors.HexColor("#020817")
NAVY_2 = colors.HexColor("#061B2E")
NAVY_3 = colors.HexColor("#0D2A3F")
TEAL = colors.HexColor("#00A7B5")
TEAL_DARK = colors.HexColor("#007481")
CYAN = colors.HexColor("#5FE7F2")
MONASH_BLUE = colors.HexColor("#006DAE")
INK = colors.HexColor("#101828")
BODY = colors.HexColor("#26364A")
MUTED = colors.HexColor("#5B6675")
LINE = colors.HexColor("#D4E3EA")
PALE = colors.HexColor("#EAFBFC")
SOFT = colors.HexColor("#F6FBFC")
WHITE = colors.white


PROBLEM_STATEMENT = (
    "Collection crews still follow fixed routes, causing them to empty underfilled bins while high-use sites "
    "overflow. This wastes fuel, increases collection emissions, and leaves operators without real-time visibility "
    "into which bins need urgent service."
)

PROPOSED_SOLUTION = (
    "BinSight is a street-block prototype with three smart bins, one return station and hub. Focus A instruments "
    "each bin with a Teensy 4.1 running FreeRTOS for deterministic ultrasonic/load-cell polling, LED states, "
    "watchdog recovery and confidence flags for blocked or noisy readings. Focus C uses a Raspberry Pi/laptop "
    "hub to log sensor and QR events, estimate time-to-overflow with a tree-based model, rank pickups and "
    "compare fixed schedules against priority routes. Focus B adds a camera classifier that accepts plastic, "
    "metal and glass, then rejects non-recyclables before they enter the stream. Focus D builds an ESP32 return "
    "station following Singapore's Beverage Container Return Scheme / Return Right RVM flow with QR session, "
    "chute/servo feedback, dashboard logs and a simulated refund."
)

VALIDATION_SUMMARY = (
    "The simulation uses prototype logs in a 30-day district model. It compares fixed collection with the priority "
    "route and checks for fewer overflows, fewer wasted trips, shorter route distance, lower fuel and CO2, cleaner "
    "return-station streams, and lower sensing energy within the USD150/SGD200 build."
)

def require_logo() -> Path:
    if not LOGO_PNG.exists():
        raise FileNotFoundError(f"Missing Monash logo asset: {LOGO_PNG}")
    return LOGO_PNG


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.replace("#", ""))


def styles():
    return {
        "body": ParagraphStyle(
            "body",
            fontName="Helvetica",
            fontSize=10.0,
            leading=13.0,
            textColor=BODY,
            alignment=TA_LEFT,
        ),
        "small": ParagraphStyle(
            "small",
            fontName="Helvetica",
            fontSize=7.25,
            leading=9.15,
            textColor=BODY,
            alignment=TA_LEFT,
        ),
        "tiny": ParagraphStyle(
            "tiny",
            fontName="Helvetica",
            fontSize=6.85,
            leading=8.55,
            textColor=BODY,
            alignment=TA_LEFT,
        ),
        "white": ParagraphStyle(
            "white",
            fontName="Helvetica",
            fontSize=7.35,
            leading=9.2,
            textColor=colors.HexColor("#EAFBFC"),
            alignment=TA_LEFT,
        ),
        "cover": ParagraphStyle(
            "cover",
            fontName="Helvetica",
            fontSize=8.1,
            leading=9.7,
            textColor=colors.HexColor("#D8F6FA"),
            alignment=TA_CENTER,
        ),
    }


def para(c, text, style, x, y_top, w, h=1000):
    p = Paragraph(text, style)
    _, used_h = p.wrap(w, h)
    p.drawOn(c, x, y_top - used_h)
    return used_h


def draw_logo(c, logo_path, x, y, width=43 * mm, height=12.3 * mm):
    c.drawImage(str(logo_path), x, y, width=width, height=height, mask="auto")


def draw_header(c, logo_path, subtitle):
    c.saveState()
    c.setFillColor(NAVY)
    c.rect(0, PAGE_H - 26.5 * mm, PAGE_W, 26.5 * mm, fill=1, stroke=0)
    c.setFillColor(TEAL)
    c.rect(0, PAGE_H - 26.5 * mm, PAGE_W, 1.2 * mm, fill=1, stroke=0)
    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 10.5)
    c.drawString(16 * mm, PAGE_H - 13 * mm, "BinSight")
    c.setFillColor(colors.HexColor("#B8DCE4"))
    c.setFont("Helvetica", 7.35)
    c.drawString(16 * mm, PAGE_H - 19.5 * mm, subtitle)
    draw_logo(c, logo_path, PAGE_W - 58 * mm, PAGE_H - 20.2 * mm)
    c.restoreState()


def draw_footer(c, page_num, total_pages=1):
    c.saveState()
    c.setStrokeColor(LINE)
    c.line(16 * mm, 13 * mm, PAGE_W - 16 * mm, 13 * mm)
    c.setFillColor(MUTED)
    c.setFont("Helvetica", 7.25)
    c.drawString(16 * mm, 8 * mm, "Southeast Asia Engineering Design Competition 2026 | Team: MON BLUE")
    c.drawRightString(PAGE_W - 16 * mm, 8 * mm, f"Page {page_num} of {total_pages}")
    c.restoreState()


def panel(c, x, y_top, w, h, title, accent=TEAL, fill=WHITE, title_color=INK):
    c.saveState()
    c.setFillColor(fill)
    c.setStrokeColor(LINE)
    c.roundRect(x, y_top - h, w, h, 7, fill=1, stroke=1)
    c.setFillColor(accent)
    c.roundRect(x, y_top - 4.5, w, 4.5, 3, fill=1, stroke=0)
    c.setFillColor(title_color)
    c.setFont("Helvetica-Bold", 10.6)
    c.drawString(x + 9, y_top - 18.5, title)
    c.restoreState()
    return x + 9, y_top - 29, w - 18


def bullets(c, items, style, x, y_top, w, color=TEAL, gap=2.9):
    y = y_top
    for item in items:
        c.setFillColor(color)
        c.circle(x + 2.1, y - 3.5, 1.45, fill=1, stroke=0)
        used = para(c, item, style, x + 8, y, w - 8, 80)
        y -= used + gap
    return y


def section_block(c, title, text, style, x, y_top, w, accent=TEAL):
    c.saveState()
    c.setFillColor(INK)
    c.setFont("Helvetica-Bold", 12.3)
    c.drawString(x, y_top, title)
    rule_y = y_top - 5.0 * mm
    c.setStrokeColor(LINE)
    c.setLineWidth(0.6)
    c.line(x, rule_y, x + w, rule_y)
    c.setStrokeColor(accent)
    c.setLineWidth(1.6)
    c.line(x, rule_y, x + 39 * mm, rule_y)
    c.restoreState()
    used = para(c, text, style, x, y_top - 12 * mm, w, 80 * mm)
    return y_top - 12 * mm - used - 15 * mm


def cover_page(c, logo_path):
    c.saveState()
    c.setFillColor(NAVY)
    c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)

    draw_logo(c, logo_path, PAGE_W - 62 * mm, PAGE_H - 25 * mm, 47 * mm, 13.4 * mm)

    c.setStrokeColor(TEAL)
    c.setLineWidth(2.6)
    c.line(24 * mm, PAGE_H - 46 * mm, 72 * mm, PAGE_H - 46 * mm)

    c.setFillColor(colors.HexColor("#B9DCE4"))
    c.setFont("Helvetica-Bold", 10.2)
    c.drawString(24 * mm, PAGE_H - 65 * mm, "SOUTHEAST ASIA ENGINEERING DESIGN COMPETITION 2026")
    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 58)
    c.drawString(23.5 * mm, PAGE_H - 96 * mm, "BinSight")
    c.setFillColor(CYAN)
    c.setFont("Helvetica-Bold", 15.3)
    c.drawString(25 * mm, PAGE_H - 109 * mm, "Smart waste sensing, recycling-return validation")
    c.drawString(25 * mm, PAGE_H - 118 * mm, "and route simulation")
    c.setFillColor(colors.HexColor("#E8FBFC"))
    c.setFont("Helvetica-Bold", 12.5)
    c.drawString(25 * mm, PAGE_H - 139 * mm, "Team: MON BLUE")

    c.setStrokeColor(colors.HexColor("#1E6F7A"))
    c.setLineWidth(0.8)
    c.line(24 * mm, 58 * mm, PAGE_W - 24 * mm, 58 * mm)
    c.setFillColor(colors.HexColor("#CDEEF3"))
    c.setFont("Helvetica-Bold", 9.0)
    c.drawString(25 * mm, 47 * mm, "Three sensor bins | QR return station | Local decision hub | Fixed-vs-priority route simulation")
    c.restoreState()


def content_page_one(c, logo_path):
    st = styles()
    c.setFillColor(WHITE)
    c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    draw_header(c, logo_path, "Team: MON BLUE")

    left = 22 * mm
    right = PAGE_W - 22 * mm
    full_w = right - left
    top = PAGE_H - 41 * mm

    c.setFillColor(INK)
    c.setFont("Helvetica-Bold", 20.0)
    c.drawString(left, top, "Engineering Proposal")

    y = top - 19 * mm
    y = section_block(c, "Problem Statement", PROBLEM_STATEMENT, st["body"], left, y, full_w, TEAL)
    y = section_block(c, "Proposed Solution", PROPOSED_SOLUTION, st["body"], left, y, full_w, MONASH_BLUE)
    section_block(c, "Simulation / Proposed Targets", VALIDATION_SUMMARY, st["body"], left, y, full_w, TEAL)

    draw_footer(c, 1, 1)


def build_pdf():
    logo_path = require_logo()
    c = canvas.Canvas(str(PDF_PATH), pagesize=A4)
    c.setTitle("BinSight Engineering Proposal")
    c.setAuthor("MON BLUE")
    c.setSubject("Southeast Asia Engineering Design Competition 2026")
    cover_page(c, logo_path)
    c.showPage()
    content_page_one(c, logo_path)
    c.showPage()
    c.save()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="101828", size=8.1):
    cell.text = ""
    pgh = cell.paragraphs[0]
    pgh.paragraph_format.space_after = Pt(0)
    run = pgh.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_header(section, logo_path, label):
    section.header_distance = Cm(0.45)
    header = section.header
    header.paragraphs[0].text = ""
    table = header.add_table(rows=1, cols=2, width=Cm(18.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(9.5)
    table.columns[1].width = Cm(8.5)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "020817")
    left = table.rows[0].cells[0].paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run(f"BinSight | {label}")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb("FFFFFF")
    right = table.rows[0].cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    run = right.add_run()
    run.add_picture(str(logo_path), width=Cm(4.3))


def add_paragraph_bottom_border(pgh, color="00A7B5"):
    p_pr = pgh._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def docx_para(doc, text, size=10.0, color="101828", bold=False, after=5):
    pgh = doc.add_paragraph()
    pgh.paragraph_format.space_after = Pt(after)
    pgh.paragraph_format.line_spacing = 1.08
    run = pgh.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    return pgh


def docx_heading(doc, text, size=12.5, color="006DAE", rule=True):
    pgh = doc.add_paragraph()
    pgh.paragraph_format.space_before = Pt(11)
    pgh.paragraph_format.space_after = Pt(6)
    run = pgh.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if rule:
        add_paragraph_bottom_border(pgh, color)


def add_table_docx(doc, rows, widths=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Cm(width)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if r_idx == 0:
                set_cell_shading(cell, "061B2E")
                set_cell_text(cell, rows[r_idx][c_idx], True, "FFFFFF", 7.8)
            else:
                set_cell_shading(cell, "F6FBFC" if r_idx % 2 else "FFFFFF")
                set_cell_text(cell, rows[r_idx][c_idx], False, "101828", 7.3)
    return table


def build_docx():
    logo_path = require_logo()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    add_docx_header(section, logo_path, "Team: MON BLUE")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(104)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("BinSight")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(42)
    run.font.color.rgb = rgb("020817")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Smart waste sensing, recycling-return validation and route simulation")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(14)
    run.font.color.rgb = rgb("00A7B5")

    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team.add_run("Team: MON BLUE")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(12)
    run.font.color.rgb = rgb("061B2E")

    scope = docx_para(
        doc,
        "Three sensor bins | QR return station | Local decision hub | Fixed-vs-priority route simulation",
        9.4,
        "344054",
        True,
        0,
    )
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    section2.page_width = Cm(21)
    section2.page_height = Cm(29.7)
    section2.top_margin = Cm(1.35)
    section2.bottom_margin = Cm(1.0)
    section2.left_margin = Cm(1.55)
    section2.right_margin = Cm(1.55)
    add_docx_header(section2, logo_path, "Page 1 of 1")

    docx_heading(doc, "Engineering Proposal", 16, "101828", False)
    docx_heading(doc, "Problem Statement", 12.2, "007481")
    docx_para(doc, PROBLEM_STATEMENT)
    docx_heading(doc, "Proposed Solution", 12.2, "006DAE")
    docx_para(doc, PROPOSED_SOLUTION)
    docx_heading(doc, "Simulation / Proposed Targets", 12.2, "007481")
    docx_para(doc, VALIDATION_SUMMARY)

    doc.save(DOCX_PATH)


def build_all():
    build_pdf()
    build_docx()
    print(f"PDF: {PDF_PATH}")
    print(f"DOCX: {DOCX_PATH}")


if __name__ == "__main__":
    build_all()
