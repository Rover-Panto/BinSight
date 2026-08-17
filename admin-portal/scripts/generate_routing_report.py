from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "ROUTING_REPORT.md"
OUTPUT_DIR = ROOT / "reports"
OUTPUT_DOCX = OUTPUT_DIR / "BinSight_Routing_Subsystem_Report_Improved.docx"

# Resolved design system: standard_business_brief.
# Page: Letter portrait, 1.0-inch margins, 0.492-inch header/footer distance.
# Body: Calibri 11 pt, 6 pt after, 1.10 line spacing.
# H1/H2/H3: 16/13/12 pt, prescribed blue palette and spacing.
# Lists: 0.25-inch marker, 0.5-inch text, 0.25-inch hanging indent,
#        8 pt after, 1.167 line spacing.
# Tables: 9360 DXA, 120 DXA indent, 80/80/120/120 DXA cell margins,
#         single grid, F2F4F7 header fill.

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "68717A"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
TEAL = "0F766E"
WHITE = "FFFFFF"
BLACK = "1A1A1A"
TABLE_WIDTH_DXA = 9360


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("left", start),
        ("bottom", bottom),
        ("right", end),
    ):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def insert_paragraph_shading(p_pr, shd):
    earlier_tags = {
        qn("w:pStyle"),
        qn("w:keepNext"),
        qn("w:keepLines"),
        qn("w:pageBreakBefore"),
        qn("w:framePr"),
        qn("w:widowControl"),
        qn("w:numPr"),
        qn("w:suppressLineNumbers"),
        qn("w:pBdr"),
    }
    insert_at = 0
    for index, child in enumerate(p_pr):
        if child.tag in earlier_tags:
            insert_at = index + 1
    p_pr.insert(insert_at, shd)


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        later_tags = {
            qn("w:tblBorders"),
            qn("w:shd"),
            qn("w:tblLayout"),
            qn("w:tblCellMar"),
            qn("w:tblLook"),
        }
        insert_at = next(
            (idx for idx, child in enumerate(tbl_pr) if child.tag in later_tags),
            len(tbl_pr),
        )
        tbl_pr.insert(insert_at, tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([r_pr, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(https?://\S+|\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def add_rich_text(paragraph, text):
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, color=BLACK)
        token = match.group(0)
        if token.startswith("http"):
            url = token.rstrip(".,)")
            suffix = token[len(url):]
            add_hyperlink(paragraph, url, url)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_font(run, color=BLACK)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, color=BLACK, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, color=BLACK, italic=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, color=BLACK)


def add_numbering(document):
    numbering = document.part.numbering_part.element
    used_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    next_abs = max(used_abs, default=-1) + 1
    used_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_num = max(used_num, default=0) + 1

    def make_abstract(abstract_id, fmt, text_value, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, spacing, ind])
        level.extend([start, num_fmt, suff, lvl_text, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            level.append(r_pr)
        abstract.append(level)
        return abstract

    def make_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        return num

    # OOXML requires every abstractNum before every num. Insert new abstract
    # definitions before the first existing num so Word does not repair/drop them.
    first_num_index = next(
        (idx for idx, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, make_abstract(next_abs, "bullet", "•", "Calibri"))
    numbering.insert(first_num_index + 1, make_abstract(next_abs + 1, "decimal", "%1."))
    numbering.append(make_num(next_num, next_abs))
    numbering.append(make_num(next_num + 1, next_abs + 1))
    return next_num, next_num + 1, next_abs, next_abs + 1


def add_num_instance(document, abstract_id):
    numbering = document.part.numbering_part.element
    used_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(used_num, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def set_paragraph_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
    else:
        p_pr.remove(num_pr)
        for child in list(num_pr):
            num_pr.remove(child)
    p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    zoom = document.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = document.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("BinSight  |  Predictive Collection Routing")
    set_run_font(r, size=9, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_cover(document):
    for _ in range(5):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    r = kicker.add_run("TECHNICAL IMPLEMENTATION REPORT  |  FOCUS AREA C")
    set_run_font(r, size=10, color=TEAL, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("BinSight Predictive")
    set_run_font(r, size=28, color=NAVY, bold=True)
    r.add_break()
    r = title.add_run("Collection Routing Subsystem")
    set_run_font(r, size=28, color=NAVY, bold=True)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(44)
    r = sub.add_run("OpenStreetMap-based routing for a Subang Jaya pilot")
    set_run_font(r, size=14, color=DARK_BLUE)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(22)
    add_rich_text(meta, "Team MON BLUE  |  Southeast Asia Engineering Design Competition 2026")

    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status.paragraph_format.space_before = Pt(6)
    status.paragraph_format.space_after = Pt(6)
    status.paragraph_format.left_indent = Inches(0.65)
    status.paragraph_format.right_indent = Inches(0.65)
    status.paragraph_format.line_spacing = 1.15
    p_pr = status._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    insert_paragraph_shading(p_pr, shd)
    r = status.add_run("SIMULATION LOCKED  |  FIELD VALIDATION PENDING")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_before = Pt(42)
    r = date.add_run("August 2026")
    set_run_font(r, size=11, color=MUTED)
    document.add_page_break()


def parse_table(lines, index):
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def table_widths(headers):
    key = tuple(h.lower().replace("`", "") for h in headers)
    if len(headers) == 2:
        return [2850, 6510]
    if len(headers) == 3:
        return [2300, 2200, 4860]
    if len(headers) == 4 and "interpretation" in key:
        return [2200, 1700, 1900, 3560]
    if len(headers) == 4 and "area" in key:
        return [1100, 3400, 1800, 3060]
    if len(headers) == 4:
        return [1800, 2600, 1800, 3160]
    base = TABLE_WIDTH_DXA // len(headers)
    widths = [base] * len(headers)
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def add_table(document, rows):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = table_widths(rows[0])
    for row_idx, row in enumerate(rows):
        prevent_row_split(table.rows[row_idx])
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_rich_text(p, value)
            for run in p.runs:
                set_run_font(run, size=9.2 if len(rows[0]) >= 4 else 9.5, color=BLACK, bold=row_idx == 0)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
        if row_idx == 0:
            set_repeat_table_header(table.rows[row_idx])
    set_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def add_code_block(document, code_lines):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    insert_paragraph_shading(p_pr, shd)
    for idx, line in enumerate(code_lines):
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color=DARK_BLUE)
        if idx < len(code_lines) - 1:
            run.add_break()


def add_callout(document, text):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.22)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    insert_paragraph_shading(p_pr, shd)
    add_rich_text(p, text)


def add_image(document, relative_path, alt_text):
    path = ROOT / relative_path
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.15))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//" + qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", alt_text)
            doc_pr.set("title", "Representative BinSight route map")


def build_document():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    document = Document()
    configure_document(document)
    bullet_num_id, decimal_num_id, bullet_abs_id, decimal_abs_id = add_numbering(document)
    add_cover(document)

    # The first markdown title and metadata are represented by the cover.
    start = next(i for i, line in enumerate(lines) if line.strip() == "## Executive summary")
    index = start
    in_code = False
    code_lines = []
    active_list_kind = None
    active_decimal_num_id = decimal_num_id
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            active_list_kind = None
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(document, code_lines)
                in_code = False
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            active_list_kind = None
            index += 1
            continue
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            # Preserve machine-readable result-lock markers in Markdown without
            # exposing them in the competition-facing Word/PDF report.
            active_list_kind = None
            index += 1
            continue
        if stripped.startswith("!["):
            active_list_kind = None
            match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            if match:
                add_image(document, match.group(2), match.group(1))
            index += 1
            continue
        if stripped.startswith("|"):
            active_list_kind = None
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue
        if stripped.startswith("### "):
            active_list_kind = None
            p = document.add_paragraph(style="Heading 3")
            add_rich_text(p, stripped[4:])
            index += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:]
            active_list_kind = None
            if heading == "Appendix A. Site-to-controller schedule":
                document.add_page_break()
            p = document.add_paragraph(style="Heading 2")
            add_rich_text(p, heading)
            index += 1
            continue
        if stripped.startswith("# "):
            active_list_kind = None
            p = document.add_paragraph(style="Heading 1")
            add_rich_text(p, stripped[2:])
            index += 1
            continue
        if stripped.startswith("> "):
            active_list_kind = None
            add_callout(document, stripped[2:])
            index += 1
            continue
        if re.match(r"^- ", stripped):
            active_list_kind = "bullet"
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            set_paragraph_numbering(p, bullet_num_id)
            add_rich_text(p, stripped[2:])
            index += 1
            continue
        number_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if number_match:
            if active_list_kind != "number":
                active_decimal_num_id = add_num_instance(document, decimal_abs_id)
            active_list_kind = "number"
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            set_paragraph_numbering(p, active_decimal_num_id)
            add_rich_text(p, number_match.group(1))
            index += 1
            continue

        active_list_kind = None
        p = document.add_paragraph()
        if stripped.startswith("**Figure"):
            p.style = document.styles["Caption"]
        add_rich_text(p, stripped)
        index += 1

    # Keep heading structure and page furniture stable in Word and LibreOffice.
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True

    core_props = document.core_properties
    core_props.title = "BinSight Predictive Collection Routing Subsystem"
    core_props.subject = "Focus Area C routing implementation report"
    core_props.author = "Team MON BLUE"
    core_props.keywords = "BinSight, routing, OpenStreetMap, OSRM, OR-Tools, Subang Jaya"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
